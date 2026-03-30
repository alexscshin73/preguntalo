from dataclasses import dataclass
from io import BytesIO
import re

from docx import Document
from pypdf import PdfReader


@dataclass
class ParsedPage:
    page_number: int
    text: str


@dataclass
class ParsedLine:
    page_number: int
    text: str
    role: str = "body"


@dataclass
class ParsedSection:
    heading: str
    body_lines: list[ParsedLine]
    order: int
    page_start: int
    page_end: int
    language: str = "en"

    @property
    def body_text(self) -> str:
        return "\n".join(line.text for line in self.body_lines).strip()


@dataclass
class ParsedDocument:
    pages: list[ParsedPage]
    sections: list[ParsedSection]


class DocumentParser:
    @staticmethod
    def supported_extensions() -> set[str]:
        return {"pdf", "docx", "txt", "md"}

    def parse(self, filename: str, content: bytes) -> ParsedDocument:
        extension = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
        lines, pages = self._extract_lines(extension, content)
        sections = self._build_sections(lines)
        return ParsedDocument(pages=pages, sections=sections)

    def _extract_lines(self, extension: str, content: bytes) -> tuple[list[ParsedLine], list[ParsedPage]]:
        if extension in {"txt", "md"}:
            text = content.decode("utf-8", errors="ignore")
            return self._extract_stream_document(text)

        if extension == "docx":
            document = Document(BytesIO(content))
            text = "\n".join((paragraph.text or "").strip() for paragraph in document.paragraphs)
            return self._extract_stream_document(text)

        if extension == "pdf":
            reader = PdfReader(BytesIO(content))
            pages: list[ParsedPage] = []
            page_line_buckets: list[tuple[int, list[str]]] = []
            for index, page in enumerate(reader.pages, start=1):
                text = (page.extract_text() or "").strip()
                if not text:
                    text = f"[OCR required] Page {index} has no extractable text."
                pages.append(ParsedPage(page_number=index, text=text))
                normalized_lines = [
                    stripped
                    for raw_line in text.splitlines()
                    if (stripped := self._normalize_line_text(raw_line))
                ]
                page_line_buckets.append((index, normalized_lines))

            repeated_lines = self._repeated_pdf_lines(page_line_buckets)
            lines: list[ParsedLine] = []
            for page_number, normalized_lines in page_line_buckets:
                for stripped in normalized_lines:
                    if stripped in repeated_lines:
                        continue
                    lines.append(ParsedLine(page_number=page_number, text=stripped, role=self._classify_line(stripped)))
            return lines, pages

        raise ValueError(f"Unsupported file type: .{extension or 'unknown'}")

    def _extract_stream_document(self, text: str) -> tuple[list[ParsedLine], list[ParsedPage]]:
        normalized_text = text.replace("\r\n", "\n").replace("\r", "\n")
        raw_pages = [segment for segment in normalized_text.split("\f")]
        pages: list[ParsedPage] = []
        lines: list[ParsedLine] = []

        for page_index, raw_page in enumerate(raw_pages, start=1):
            page_number = page_index
            page_lines: list[str] = []

            for raw_line in raw_page.splitlines():
                stripped = self._normalize_line_text(raw_line)
                if not stripped:
                    continue

                page_marker_number = self._extract_page_marker_number(stripped)
                if page_marker_number is not None and not page_lines:
                    page_number = page_marker_number

                page_lines.append(stripped)
                lines.append(
                    ParsedLine(
                        page_number=page_number,
                        text=stripped,
                        role=self._classify_line(stripped),
                    )
                )

            pages.append(ParsedPage(page_number=page_number, text="\n".join(page_lines).strip()))

        if not pages:
            pages.append(ParsedPage(page_number=1, text=""))

        return lines, pages

    def _build_sections(self, lines: list[ParsedLine]) -> list[ParsedSection]:
        cleaned = [line for line in lines if line.text.strip()]
        if not cleaned:
            return [
                ParsedSection(
                    heading="Imported content",
                    body_lines=[ParsedLine(page_number=1, text="No readable text found.", role="body")],
                    order=0,
                    page_start=1,
                    page_end=1,
                )
            ]

        sections: list[ParsedSection] = []
        current_heading = "Imported content"
        current_body: list[ParsedLine] = []
        current_page_start = cleaned[0].page_number
        current_page_end = cleaned[0].page_number

        def flush_section() -> None:
            nonlocal current_body, current_heading, current_page_start, current_page_end
            if not current_body:
                return
            sections.append(
                ParsedSection(
                    heading=current_heading,
                    body_lines=current_body.copy(),
                    order=len(sections),
                    page_start=current_page_start,
                    page_end=current_page_end,
                    language=self._detect_language(" ".join(line.text for line in current_body)),
                )
            )
            current_body.clear()

        for line in cleaned:
            if line.role == "heading":
                flush_section()
                current_heading = line.text
                current_page_start = line.page_number
                current_page_end = line.page_number
                continue

            if line.role == "page_marker":
                current_page_end = max(current_page_end, line.page_number)
                continue

            current_body.append(line)
            current_page_end = line.page_number

        flush_section()

        if not sections:
            first_line = cleaned[0]
            aggregated = [
                ParsedLine(page_number=line.page_number, text=line.text, role=line.role) for line in cleaned
            ]
            sections.append(
                ParsedSection(
                    heading=first_line.text,
                    body_lines=aggregated,
                    order=0,
                    page_start=cleaned[0].page_number,
                    page_end=cleaned[-1].page_number,
                    language=self._detect_language(" ".join(line.text for line in aggregated)),
                )
            )

        return sections

    def _classify_line(self, text: str) -> str:
        if not text:
            return "body"
        normalized = text.strip()
        lower = normalized.lower()
        if self._looks_like_page_marker(lower):
            return "page_marker"
        if normalized.startswith("[OCR required]"):
            return "ocr_hint"
        if self._looks_like_heading(normalized):
            return "heading"
        if self._looks_like_image_hint(normalized, lower):
            return "caption"
        return "body"

    def _looks_like_heading(self, line: str) -> bool:
        if line.startswith("![") or line.startswith("[OCR required]"):
            return False
        if re.fullmatch(r"\d+(\.\d+)*", line):
            return False
        if len(line) > 90:
            return False
        if re.match(r"^#{1,6}\s+\S+", line):
            return True
        if line.endswith(":"):
            return True
        tokens = line.split()
        if len(tokens) <= 6 and len(line) >= 6 and not re.search(r"[.!?]$", line):
            return True
        if re.match(r"^\d+(\.\d+)*\s", line):
            return True
        if line.isupper() and len(line) >= 3:
            return True
        return False

    def _looks_like_page_marker(self, lower_line: str) -> bool:
        if re.fullmatch(r"\d+(\.\d+)*", lower_line):
            return True
        return bool(re.match(r"^(page|pagina|página|pagina|페이지)\s*\d+", lower_line))

    def _extract_page_marker_number(self, line: str) -> int | None:
        match = re.match(r"^(?:page|pagina|página|페이지)\s*(\d+)", line.lower())
        if match is None:
            return None
        return int(match.group(1))

    def _looks_like_image_hint(self, line: str, lower_line: str) -> bool:
        if re.match(r"^(figure|figura|imagen|image|그림|도표|이미지)\b", lower_line):
            return True
        if re.match(r"^!\[[^\]]*\]\([^)]+\)$", line):
            return True
        return bool(re.search(r"\.(png|jpg|jpeg|gif|webp|svg)$", lower_line))

    def _normalize_line_text(self, text: str) -> str:
        stripped = text.strip()
        if not stripped:
            return ""

        markdown_image = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)$", stripped)
        if markdown_image:
            alt_text = markdown_image.group(1).strip() or markdown_image.group(2).strip()
            return f"Image: {alt_text}"

        stripped = re.sub(r"^#{1,6}\s+", "", stripped)
        stripped = re.sub(r"\s+", " ", stripped).strip()
        if re.search(r"[가-힣A-Za-zÁÉÍÓÚáéíóúÑñÜü)\]]\s+\d+$", stripped) and not re.search(r"[.!?:]$", stripped):
            stripped = re.sub(r"\s+\d+$", "", stripped).strip()
        return stripped

    def _repeated_pdf_lines(self, page_line_buckets: list[tuple[int, list[str]]]) -> set[str]:
        occurrences: dict[str, set[int]] = {}

        for page_number, lines in page_line_buckets:
            seen_on_page: set[str] = set()
            for line in lines:
                normalized = self._pdf_repeat_key(line)
                if not normalized or normalized in seen_on_page:
                    continue
                seen_on_page.add(normalized)
                occurrences.setdefault(normalized, set()).add(page_number)

        repeated_keys = {
            key
            for key, pages in occurrences.items()
            if len(pages) >= 2 and len(key) <= 80
        }

        repeated_lines: set[str] = set()
        for _, lines in page_line_buckets:
            for line in lines:
                if self._pdf_repeat_key(line) in repeated_keys:
                    repeated_lines.add(line)

        return repeated_lines

    def _pdf_repeat_key(self, line: str) -> str:
        normalized = re.sub(r"\s+", " ", line).strip().lower()
        normalized = re.sub(r"\s+\d+$", "", normalized)
        normalized = normalized.strip()
        if len(normalized) < 4:
            return ""
        return normalized

    def _detect_language(self, text: str) -> str:
        if not text:
            return "en"
        if re.search(r"[ㄱ-ㅎ가-힣]", text):
            return "ko"
        if re.search(r"[áéíóúñüÁÉÍÓÚÑÜ¿¡]", text):
            return "es"
        return "en"
